from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import fitz  # PyMuPDF
import re
import io
import zipfile
import pandas as pd
from docx import Document

app = Flask(__name__)
CORS(app) # للسماح لموقعك بالاتصال بهذا الـ API

email_pattern = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
phone_pattern = re.compile(r"(?<!\d)(?:(?:\+|00)\d{1,3}[\s\-.]?|0\d{1,2}[\s\-.]?)[\d\s\-]{6,10}\d(?!\d)")
linkedin_pattern = re.compile(r"linkedin\.com/in/[a-zA-Z0-9_-]+")
patterns = [email_pattern, phone_pattern, linkedin_pattern]

@app.route('/', methods=['GET'])
def home():
    return jsonify({"status": "API is running. Send POST request with files to /redact"})

@app.route('/redact', methods=['POST'])
def redact_cvs():
    if 'files' not in request.files:
        return jsonify({"error": "No files uploaded"}), 400

    uploaded_files = request.files.getlist('files')
    all_candidates_data = []
    
    try:
        memory_zip = io.BytesIO()
        with zipfile.ZipFile(memory_zip, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for uploaded_file in uploaded_files:
                file_ext = uploaded_file.filename.split('.')[-1].lower()
                output_filename = f"REDACTED_{uploaded_file.filename}"
                full_text_for_extraction = ""
                output_buffer = io.BytesIO()
                
                # --- PDF PROCESSING ---
                if file_ext == "pdf":
                    pdf_bytes = uploaded_file.read()
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    for page in doc:
                        text = page.get_text("text")
                        full_text_for_extraction += text + "\n"
                        for pattern in patterns:
                            for match in pattern.finditer(text):
                                sensitive_text = match.group()
                                text_instances = page.search_for(sensitive_text)
                                for inst in text_instances:
                                    page.add_redact_annot(inst, fill=(0, 0, 0))
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)
                    doc.save(output_buffer, garbage=4, deflate=True)
                    doc.close()

                # --- DOCX PROCESSING ---
                elif file_ext == "docx":
                    docx_bytes = uploaded_file.read()
                    doc_docx = Document(io.BytesIO(docx_bytes))
                    full_text_for_extraction = "\n".join([para.text for para in doc_docx.paragraphs])
                    
                    def replace_text_in_run(run):
                        for pattern in patterns:
                            if pattern.search(run.text):
                                run.text = pattern.sub("[REDACTED]", run.text)

                    for para in doc_docx.paragraphs:
                        for run in para.runs:
                            replace_text_in_run(run)
                    for table in doc_docx.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        replace_text_in_run(run)
                    doc_docx.save(output_buffer)

                # --- EXCEL DATA EXTRACTION ---
                found_emails = email_pattern.findall(full_text_for_extraction)
                found_phones = phone_pattern.findall(full_text_for_extraction)
                text_lines = [line.strip() for line in full_text_for_extraction.split('\n') if line.strip()]
                guessed_name = text_lines[0] if text_lines else "Review Manually"

                all_candidates_data.append({
                    "File Name": uploaded_file.filename,
                    "Name": guessed_name,
                    "Email": found_emails[0] if found_emails else "Not Found",
                    "Phone": found_phones[0] if found_phones else "Not Found",
                })

                zip_file.writestr(output_filename, output_buffer.getvalue())

            # --- GENERATE EXCEL AND ADD TO ZIP ---
            df = pd.DataFrame(all_candidates_data)
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Candidates')
            
            zip_file.writestr("Candidate_Summary_Data.xlsx", excel_buffer.getvalue())

        memory_zip.seek(0)
        return send_file(
            memory_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name='Processed_CVs.zip'
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# هذا السطر ضروري لـ Vercel
if __name__ == '__main__':
    app.run(debug=True)
